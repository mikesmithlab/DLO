from docx import Document
import dateutil.parser
import datetime
import pathlib
import shutil
import uuid
import os
from filehandling import BatchProcess

def parse_date(datestring):
    return dateutil.parser.parse(datestring, dayfirst=True)

def process_extension():
    """This function processes the coursework extensions
    """
    not_processed=False
    for filename in BatchProcess('C:/Users/ppzmis/OneDrive - The University of Nottingham/Documents/DLO/Extensions_to_approve/*.*'):
        if process(filename=filename):
            not_processed = True
    return not_processed

def process(filename):
    filetype = pathlib.Path(filename).suffix
    path, filename = os.path.split(filename)
    if filetype == '.docx':
        manual = process_docx(filename=filename)
    else:
        manual=process_other(filename=filename, filetype=filetype)
    return manual


def process_other(filename='test.pdf', filepath='C:/Users/ppzmis/OneDrive - The University of Nottingham/Documents/DLO/', filetype='.pdf'):
    #Processes pdf and img files
    output_filename = filepath + 'Extensions_to_approve/manual/' + str(datetime.date.today().strftime("%Y_%m_%d")) + '_' + str(uuid.uuid4()) + '_' + filetype
    shutil.move(filepath+'Extensions_to_approve/'+filename, output_filename)
    return True


def process_docx(filename='test.docx', signature='signature.png', filepath ='C:/Users/ppzmis/OneDrive - The University of Nottingham/Documents/DLO/'):
    #Processes word docs
    manual=False
    doc = Document(filepath + 'Extensions_to_approve/' + filename)

    request = {
        'name':doc.tables[0].cell(0,1).text,
        'id': doc.tables[0].cell(1,1).text,
        'module': doc.tables[1].cell(1,0).text,
        'original_deadline' : parse_date(doc.tables[1].cell(1,2).text),
        'new_deadline' : parse_date(doc.tables[1].cell(1,3).text),
        }
    request['date_diff'] = (request['new_deadline']-request['original_deadline']).days

    if request['date_diff'] > 7:
        manual=True
    if 'PHYS4' in request['module']:
        manual=True

    doc.tables[1].cell(1,4).paragraphs[0].text=''
    doc.tables[1].cell(1,4).paragraphs[0].add_run().add_picture(filepath+signature)

    today = str(datetime.date.today().strftime("%d/%m/%Y"))

    for paragraph in doc.paragraphs:
        if 'Staff Signature:' in paragraph.text:
            paragraph.text = ''
            paragraph.add_run('Staff Signature:').bold = True
            paragraph.add_run('..........')
            run2=paragraph.add_run()
            run2.add_picture(filepath+signature)
            paragraph.add_run('................')
            paragraph.add_run('Staff name:').bold = True
            paragraph.add_run('........Mike Smith...........')
            paragraph.add_run('Date:').bold = True
            paragraph.add_run('..............' + today + '...................')


    if manual:
        #Move to folder for manual processing
        doc.save(filepath + 'Extensions_to_approve/manual/' + str(datetime.date.today().strftime("%Y_%m_%d")) + '_' + request['name'].replace(' ','') +'.docx')
    else:
        doc.save(filepath + 'Approved_extensions/' + str(datetime.date.today().strftime("%Y_%m_%d")) + '_' + request['name'].replace(' ','') +'.docx')

    return manual


def store_files(file_list, filepath='C:/Users/ppzmis/OneDrive - The University of Nottingham/Documents/DLO/Approved_extensions/'):
    print(file_list)
    for file in file_list:
        path, filename = os.path.split(file)
        dir_name = filepath + filename.split('_')[0] + '_' + filename.split('_')[-1].split('.')[0]
        print(dir_name)
        if not os.path.exists(dir_name):
            os.mkdir(dir_name)
        shutil.move(file, dir_name + filename)


def cleanup(filepath='C:/Users/ppzmis/OneDrive - The University of Nottingham/Documents/DLO/'):
    for file in BatchProcess(filepath+'/Extensions_to_approve/*.*'):
        os.remove(file)
    for file in BatchProcess(filepath+'/Approved_extensions/*.*'):
        os.remove(file)







