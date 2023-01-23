from docx import Document
import dateutil.parser
import datetime
import pathlib
import shutil
import uuid
import os, sys

from filehandling import BatchProcess
from pydates.pydates import parse_date, relative_datetime, format_datetime_to_str

parent = os.path.abspath('.')
sys.path.insert(1, parent)

from addresses import DLO_DIR


def process_extension():
    """This function processes the coursework extensions
    If function returns True further manual processing required
    """
    manual_processing=False
    for filename in BatchProcess(DLO_DIR + 'Extensions_to_approve/*.*'):
        if process(filename=filename):
            manual_processing = True
    return manual_processing

def process(filename):
    filetype = pathlib.Path(filename).suffix
    path, filename = os.path.split(filename)

    if filetype == '.docx':
        manual = process_docx(filename=filename)
    else:
        manual=process_other(filename=filename, filetype=filetype)
    return manual


def process_other(filename='test.pdf', filepath=DLO_DIR, filetype='.pdf'):
    #Processes pdf and img files
    output_filename = filepath + 'Extensions_to_approve/manual/' + str(datetime.date.today().strftime("%Y_%m_%d")) + '_' + str(uuid.uuid4()) + '_' + filetype
    shutil.move(filepath+'Extensions_to_approve/'+filename, output_filename)
    return True


def process_docx(filename='test.docx', signature='signature.png', filepath = DLO_DIR):
    #Processes word docs
    manual=False
    doc = Document(filepath + 'Extensions_to_approve/' + filename)

    if True:

        request = {
            'name':doc.tables[0].cell(0,1).text,
            'id': doc.tables[0].cell(1,1).text,
            'module': doc.tables[1].cell(1,0).text,
            'original_deadline' : parse_date(doc.tables[1].cell(1,2).text),
            }

        if len(doc.tables[1].cell(1,3).text) > 1:
            request['new_deadline'] : parse_date(doc.tables[1].cell(1,3).text)
        else:
            request['new_deadline'] = relative_datetime(request['original_deadline'],delta_day=7)
            doc.tables[1].cell(1,3).paragraphs[0].text = format_datetime_to_str(request['new_deadline'],format="%d/%m/%Y")

        request['date_diff'] = (request['new_deadline']-request['original_deadline']).days

        #Conditions :  if manual is still False after these the coursework can be automatically processed
        if (request['original_deadline'] -  datetime.datetime.now()).days < 0:
            #Asked for extension after original deadline
            manual=True
        elif request['date_diff'] > 7:
            #Asked for longer than 7 days
            manual=True
        elif 'PHYS4' in request['module']:
            #4th year module requires manual intervention
            manual=True

        doc.tables[1].cell(1,4).paragraphs[0].text=''
        doc.tables[1].cell(1,4).paragraphs[0].add_run().add_picture(filepath+signature)

        today = str(datetime.date.today().strftime("%d/%m/%Y"))

        for paragraph in doc.paragraphs:
            if 'Staff Signature:' in paragraph.text:
                paragraph.text = ''
                paragraph.add_run('Staff Signature:').bold = True
                paragraph.add_run('..........')
                run2=paragraph.add_run()
                run2.add_picture(filepath+signature)
                paragraph.add_run('................')
                paragraph.add_run('Staff name:').bold = True
                paragraph.add_run('........Mike Smith...........')
                paragraph.add_run('Date:').bold = True
                paragraph.add_run('..............' + today + '...................')

        print(manual)
        if manual:
            #Move to folder for manual processing
            doc.save(filepath + 'Extensions_to_approve/manual/' + str(datetime.date.today().strftime("%Y_%m_%d")) + '_' + request['name'].replace(' ','') +'.docx')
        else:
            doc.save(filepath + 'Approved_extensions/' + str(datetime.date.today().strftime("%Y_%m_%d")) + '_' + request['name'].replace(' ','') +'.docx')

        return manual
    """except:
        #If it fails process it manually
        manual=True
        doc.save(filepath + 'Extensions_to_approve/manual/' + str(datetime.date.today().strftime("%Y_%m_%d")) + '_problem.docx')
    """

def store_files(file_list, filepath=DLO_DIR + 'Approved_extensions/'):
    for file in file_list:
        path, filename = os.path.split(file)
        dir_name = filepath + filename.split('_')[0] + '_' + filename.split('_')[-1].split('.')[0]
        print(dir_name)
        if not os.path.exists(dir_name):
            os.mkdir(dir_name)
        shutil.move(file, dir_name + '/' + filename)


def cleanup(filepath=DLO_DIR):
    for file in BatchProcess(filepath+'/Extensions_to_approve/*.*'):
        os.remove(file)
    for file in BatchProcess(filepath+'/Approved_extensions/*.*'):
        os.remove(file)







